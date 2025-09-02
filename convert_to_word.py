#!/usr/bin/env python3
"""
Script pour convertir les guides Markdown en documents Word (.docx)
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def parse_markdown_to_word(markdown_file_path, output_file_path):
    """Convertit un fichier Markdown en document Word"""
    
    print(f"🔄 Conversion de {markdown_file_path}...")
    
    # Créer un nouveau document Word
    doc = Document()
    
    # Lire le fichier Markdown
    try:
        with open(markdown_file_path, 'r', encoding='utf-8') as file:
            content = file.read()
    except Exception as e:
        print(f"❌ Erreur lors de la lecture de {markdown_file_path}: {e}")
        return
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Ignorer les lignes vides
        if not line:
            i += 1
            continue
            
        # Traitement des titres
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title_text = line.lstrip('#').strip()
            
            if level == 1:
                heading = doc.add_heading(title_text, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(title_text, level=min(level, 3))
            
        # Traitement des blocs de code
        elif line.startswith('```'):
            # Récupérer le contenu du bloc de code
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # Ajouter le bloc de code au document
            if code_lines:
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
        
        # Traitement des listes à puces
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            doc.add_paragraph(text, style='List Bullet')
        
        # Traitement des listes numérotées
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            doc.add_paragraph(text, style='List Number')
        
        # Traitement du texte en gras/italique
        elif line:
            p = doc.add_paragraph()
            
            # Traiter le formatage du texte
            text = line
            
            # Remplacer **texte** par du gras
            while '**' in text:
                start = text.find('**')
                if start != -1:
                    end = text.find('**', start + 2)
                    if end != -1:
                        # Ajouter le texte avant le gras
                        if start > 0:
                            p.add_run(text[:start])
                        
                        # Ajouter le texte en gras
                        bold_text = text[start+2:end]
                        run = p.add_run(bold_text)
                        run.bold = True
                        
                        # Continuer avec le reste
                        text = text[end+2:]
                    else:
                        break
                else:
                    break
            
            # Ajouter le reste du texte
            if text:
                p.add_run(text)
        
        i += 1
    
    # Sauvegarder le document
    try:
        doc.save(output_file_path)
        print(f"✅ Document sauvegardé : {output_file_path}")
    except Exception as e:
        print(f"❌ Erreur lors de la sauvegarde : {e}")

def convert_all_guides():
    """Convertit tous les guides Markdown en Word"""
    
    guides = [
        ("GUIDE_UTILISATEUR.md", "Guide_Utilisateur_LN_Foot.docx"),
        ("GUIDE_TECHNIQUE.md", "Guide_Technique_LN_Foot.docx"), 
        ("GUIDE_DEPLOIEMENT.md", "Guide_Deploiement_LN_Foot.docx"),
        ("API_DOCUMENTATION.md", "Documentation_API_LN_Foot.docx")
    ]
    
    print("🚀 Début de la conversion des guides en Word...")
    print("=" * 50)
    
    for md_file, word_file in guides:
        if os.path.exists(md_file):
            parse_markdown_to_word(md_file, word_file)
        else:
            print(f"⚠️  Fichier non trouvé : {md_file}")
    
    print("=" * 50)
    print("✅ Conversion terminée !")
    print("\n📁 Documents Word générés :")
    
    for _, word_file in guides:
        if os.path.exists(word_file):
            size = os.path.getsize(word_file) / 1024  # taille en KB
            print(f"   • {word_file} ({size:.1f} KB)")

if __name__ == "__main__":
    convert_all_guides()
            # Récupérer le contenu du bloc de code
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # Ajouter le bloc de code
            if code_lines:
                code_text = '\n'.join(code_lines)
                code_paragraph = doc.add_paragraph(code_text)
                # Style code simple
                for run in code_paragraph.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
        
        # Traitement des listes
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Traiter le formatage gras dans les listes
            if text.startswith('**') and '**' in text[2:]:
                end_bold = text.find('**', 2)
                bold_text = text[2:end_bold]
                rest_text = text[end_bold+2:]
                
                p = doc.add_paragraph()
                p.add_run('• ').bold = False
                p.add_run(bold_text).bold = True
                if rest_text:
                    p.add_run(rest_text).bold = False
            else:
                doc.add_paragraph('• ' + text)
        
        # Traitement des listes numérotées
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(text, style='List Number')
        
        # Traitement du texte normal
        else:
            if line:
                paragraph = doc.add_paragraph()
                add_formatted_text(paragraph, line)
        
        i += 1
    
    # Sauvegarder le document
    try:
        doc.save(output_file_path)
        print(f"✅ Document Word créé : {output_file_path}")
    except Exception as e:
        print(f"❌ Erreur lors de la sauvegarde de {output_file_path}: {e}")

def add_formatted_text(paragraph, text):
    """Ajoute du texte avec formatage basique au paragraphe"""
    
    # Traitement simple du gras **texte**
    parts = re.split(r'\*\*(.*?)\*\*', text)
    
    for i, part in enumerate(parts):
        if i % 2 == 0:  # Texte normal
            if part:
                paragraph.add_run(part)
        else:  # Texte en gras
            run = paragraph.add_run(part)
            run.bold = True

def main():
    """Fonction principale pour convertir tous les guides"""
    
    print("🔄 Conversion des guides Markdown en documents Word...")
    print("=" * 50)
    
    # Définir les fichiers à convertir
    files_to_convert = [
        ('GUIDE_UTILISATEUR.md', 'GUIDE_UTILISATEUR.docx'),
        ('GUIDE_TECHNIQUE.md', 'GUIDE_TECHNIQUE.docx'),
        ('GUIDE_DEPLOIEMENT.md', 'GUIDE_DEPLOIEMENT.docx'),
        ('API_DOCUMENTATION.md', 'API_DOCUMENTATION.docx'),
    ]
    
    success_count = 0
    
    for markdown_file, word_file in files_to_convert:
        if os.path.exists(markdown_file):
            try:
                parse_markdown_to_word(markdown_file, word_file)
                success_count += 1
            except Exception as e:
                print(f"❌ Erreur lors de la conversion de {markdown_file}: {e}")
        else:
            print(f"⚠️  Fichier non trouvé : {markdown_file}")
    
    print("=" * 50)
    print(f"✅ Conversion terminée! ({success_count}/{len(files_to_convert)} fichiers convertis)")
    
    if success_count > 0:
        print("\n📄 Documents Word créés :")
        for _, word_file in files_to_convert:
            if os.path.exists(word_file):
                size = os.path.getsize(word_file) / 1024  # Taille en KB
                print(f"   • {word_file} ({size:.1f} KB)")

if __name__ == "__main__":
    main()
